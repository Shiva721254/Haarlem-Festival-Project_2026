"""
Generate the Haarlem Festival process documentation (Word + Markdown).

Sprint goals and deliverables are anchored to the real git history (feature
branches merged via pull requests into develop). Retrospective points reflect
issues actually encountered during the build. Team-specific details (member
names, ceremony dates, meeting notes) are marked for the team to complete.

Run:  python scripts/gen_process_docs.py
"""
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(OUT_DIR, exist_ok=True)

INTRO = (
    "This document describes how the Haarlem Festival ticketing website was built: "
    "the way of working, the sprint planning and the sprint retrospectives. The "
    "application is a plain-PHP MVC website (no framework) with MySQL, running in "
    "Docker, with Stripe (test mode) for payment. Development followed an Agile/Scrum "
    "approach with short, themed sprints, each delivered through one or more feature "
    "branches that were reviewed and merged into the develop branch via pull requests."
)

WAY_OF_WORKING = [
    ("Version control & branching",
     "Git with a gitflow-style model: main (release), develop (integration) and short-lived "
     "feature/* branches. Every change went through a pull request into develop, so develop "
     "always held a working, integrated build; main was promoted from develop as a release."),
    ("Code review",
     "Each feature branch was opened as a pull request with a detailed description and merged "
     "after review. Conflicts on shared files (e.g. the layout header) were resolved by rebasing "
     "the branch onto the latest develop."),
    ("Architecture discipline",
     "A Controller -> Service -> Repository layering, with every service and repository consumed "
     "through an interface (programming against abstractions). This kept features consistent and "
     "independently reviewable."),
    ("Database changes",
     "All schema and seed changes were made as numbered, forward-only SQL migrations applied by "
     "database/migrate.php, so every environment converged to the same state."),
    ("Definition of Done",
     "A feature was 'done' when it was implemented behind interfaces, ran end-to-end against the "
     "running app (not just unit-level), passed a lint check, was committed in small reviewable "
     "commits, and merged via PR into develop."),
]

SPRINTS = [
    {
        "title": "Sprint 1 — Foundation & Authentication",
        "goal": "Stand up the project skeleton and the user/account foundation so later features "
                "have a base to build on.",
        "delivered": [
            "Docker stack (nginx, PHP-FPM, MySQL, phpMyAdmin, MailHog) and the MVC framework "
            "helpers (Container, View, Flash, base Repository, AuthMiddleware).",
            "Migration runner and the first schema (users).",
            "Registration, login/logout, email verification and password reset (MailHog).",
            "Self-service account management; role model (admin / employee / customer).",
            "Admin user management with search, sort and role filtering.",
        ],
        "well": [
            "The interface-based layering and migration system paid off immediately and were reused everywhere.",
            "Authentication, CSRF and password hashing were in from the start.",
        ],
        "improve": [
            "Branches were sometimes created from a stale local develop, causing avoidable conflicts.",
            "A real Gmail credential was briefly committed in the mail service.",
        ],
        "actions": [
            "Always sync develop (checkout + pull) before creating a feature branch.",
            "Move all secrets to .env; the credential was removed and the key revoked.",
        ],
    },
    {
        "title": "Sprint 2 — Events, catalogue & content",
        "goal": "Model the festival catalogue and present the events, and let admins manage content.",
        "delivered": [
            "Event types, events, venues, restaurants and artists with admin CRUD.",
            "The real Festival 2026 programme seeded as data (Jazz, DANCE!, Yummy, History, "
            "Magic@Teylers, Stories) with sessions, prices and capacities.",
            "Data-driven event overview and detail pages; data-driven navigation.",
            "Homepage CMS with a WYSIWYG editor and image upload.",
        ],
        "well": [
            "Treating each programme line as an event reused the whole model — no schema churn.",
            "Seeding real data early surfaced realistic pricing/VAT questions sooner.",
        ],
        "improve": [
            "An enum backing-value mismatch (role casing) and a MySQL DDL auto-commit issue broke a migration.",
        ],
        "actions": [
            "Fixed the role default via a migration and removed the transaction wrapper around DDL "
            "(MySQL auto-commits DDL).",
        ],
    },
    {
        "title": "Sprint 3 — Commerce: cart, checkout & entrance",
        "goal": "Let visitors buy tickets end-to-end and let staff validate them at the door.",
        "delivered": [
            "Shopping cart (guest + user) with live VAT-inclusive totals.",
            "Stripe (test) checkout, order creation and idempotent fulfilment.",
            "PDF tickets with QR codes and an invoice emailed via MailHog.",
            "Entrance ticket scanner for staff; admin orders list with CSV export and an order detail page.",
        ],
        "well": [
            "The cart/checkout/order/ticket flow became the backbone that passes and reservations later reused.",
        ],
        "improve": [
            "QR/PDF generation failed until the GD image extension was enabled in the PHP image.",
            "Parallel work on a teammate's repo briefly diverged the integration picture.",
        ],
        "actions": [
            "Added GD to the PHP Dockerfile; re-synced develop and verified all teammates' work was intact.",
        ],
    },
    {
        "title": "Sprint 4 — Festival features & flows",
        "goal": "Deliver the festival-specific selling features described in the brief.",
        "delivered": [
            "All-access passes (day / multi-day) for Jazz and DANCE!, reusing the ticket flow.",
            "Personal program (a customer's purchased events).",
            "Restaurant reservations: EUR 10 per-person fee + special requests (allergies), visible to admins.",
            "Pay-later orders (24h) with retry and customer order history.",
            "Stories pay-as-you-like donations and the HaarlemPas 25% reduction.",
            "Magic@Teylers kids-event page with the app-download call to action.",
        ],
        "well": [
            "Modelling passes as ticket types and discounts/donations as an effective line price avoided new plumbing.",
            "Each feature was a focused branch, kept reviews small.",
        ],
        "improve": [
            "Two older branches (personal program, admin orders) drifted behind develop and needed rebasing.",
        ],
        "actions": [
            "Rebased the lagging branches onto develop and resolved the conflicts before merging.",
        ],
    },
    {
        "title": "Sprint 5 — Experience, content, compliance & docs",
        "goal": "Polish the visitor experience, add participant content, and meet the non-functional "
                "and documentation requirements.",
        "delivered": [
            "Homepage showing every event with prices, the festival passes with real prices, a "
            "condensed day-by-day schedule and a map of locations.",
            "Participant (artist) detail pages: gallery, career highlights, tracks, a simulated audio "
            "sample and the schedule; an admin gallery-upload UI; real content sourced from "
            "Wikipedia/Wikimedia Commons.",
            "Security & GDPR pass: hardened session cookie, privacy policy, registration consent, "
            "self-service data export and account erasure (anonymisation).",
            "Technical documentation (ERD + UML class diagram) generated from the live schema.",
        ],
        "well": [
            "Building data-driven pages meant prices/schedule stay correct as the catalogue changes.",
            "Being explicit about scope let us drop a non-required feature (reCAPTCHA) to keep things minimal.",
        ],
        "improve": [
            "Some content depends on external/placeholder media that needs replacing for an offline build.",
        ],
        "actions": [
            "Added an admin gallery upload so real, licensed media can replace placeholders.",
        ],
    },
]

OVERALL = [
    "The interface-based architecture and the migration system were the two decisions that paid "
    "off most: features stayed consistent and every environment stayed reproducible.",
    "The most repeated process mistake was branching from a stale develop; the fix (sync before "
    "branching, rebase when behind) removed almost all merge pain in later sprints.",
    "Verifying features against the running application — not just at unit level — caught issues "
    "(GD extension, DDL auto-commit, pricing/VAT) that code review alone would have missed.",
    "Keeping scope tight against the brief (e.g. removing reCAPTCHA) kept the codebase minimal and focused.",
]

TODO_TEAM = [
    "Team members and their roles / lead-designer assignments per event.",
    "Actual sprint dates and the dates of the sprint ceremonies (planning, review, retrospective).",
    "Notes/screenshots from the real retrospective meetings and the sprint board (e.g. Trello/Jira).",
    "Any impediments raised with the Product Owner / teacher and how they were resolved.",
]


def build_markdown() -> None:
    lines = ["# Haarlem Festival — Process Documentation", "", "## 1. Introduction", "", INTRO, "",
             "## 2. Way of working", ""]
    for head, body in WAY_OF_WORKING:
        lines += [f"- **{head}.** {body}"]
    lines += ["", "## 3. Sprint log", ""]
    for s in SPRINTS:
        lines += [f"### {s['title']}", "", f"**Goal.** {s['goal']}", "", "**Delivered**"]
        lines += [f"- {d}" for d in s["delivered"]]
        lines += ["", "**Retrospective**", "", "| What went well | What to improve | Actions |", "|---|---|---|"]
        rows = max(len(s["well"]), len(s["improve"]), len(s["actions"]))
        for i in range(rows):
            w = s["well"][i] if i < len(s["well"]) else ""
            im = s["improve"][i] if i < len(s["improve"]) else ""
            a = s["actions"][i] if i < len(s["actions"]) else ""
            lines += [f"| {w} | {im} | {a} |"]
        lines += [""]
    lines += ["## 4. Overall retrospective", ""]
    lines += [f"- {x}" for x in OVERALL]
    lines += ["", "## 5. To complete with team-specific detail", ""]
    lines += [f"- {x}" for x in TODO_TEAM]
    out = os.path.join(OUT_DIR, "process-documentation.md")
    with open(out, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines) + "\n")
    print("wrote", out)


def build_docx() -> None:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    t = doc.add_heading("Haarlem Festival", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Process Documentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(16)
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(INTRO)

    doc.add_heading("2. Way of working", level=1)
    for head, body in WAY_OF_WORKING:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(head + ". ").bold = True
        p.add_run(body)

    doc.add_heading("3. Sprint log", level=1)
    for s in SPRINTS:
        doc.add_heading(s["title"], level=2)
        g = doc.add_paragraph()
        g.add_run("Goal. ").bold = True
        g.add_run(s["goal"])
        doc.add_paragraph("Delivered", style="Intense Quote")
        for d in s["delivered"]:
            doc.add_paragraph(d, style="List Bullet")

        doc.add_paragraph("Retrospective").runs[0].bold = True
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(("What went well", "What to improve", "Actions")):
            hdr[i].paragraphs[0].add_run(label).bold = True
        rows = max(len(s["well"]), len(s["improve"]), len(s["actions"]))
        for i in range(rows):
            cells = table.add_row().cells
            cells[0].text = s["well"][i] if i < len(s["well"]) else ""
            cells[1].text = s["improve"][i] if i < len(s["improve"]) else ""
            cells[2].text = s["actions"][i] if i < len(s["actions"]) else ""
        doc.add_paragraph("")

    doc.add_heading("4. Overall retrospective", level=1)
    for x in OVERALL:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("5. To complete with team-specific detail", level=1)
    doc.add_paragraph(
        "This draft is generated from the development history. Add the following before submission:"
    )
    for x in TODO_TEAM:
        doc.add_paragraph(x, style="List Bullet")

    out = os.path.join(OUT_DIR, "Process-Documentation.docx")
    doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    build_markdown()
    build_docx()
    print("done")
