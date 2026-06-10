"""
Generate the Haarlem Festival technical documentation (Word + Markdown) from the
real database schema and class structure. Diagrams are authored in Mermaid and
rendered to images via mermaid.ink for the Word file; the Markdown keeps the
Mermaid source so it renders on GitHub.

Run:  python scripts/gen_technical_docs.py
"""
import base64
import io
import os
import urllib.request

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(OUT_DIR, exist_ok=True)

# --------------------------------------------------------------------------- #
# Mermaid diagram sources (authored from the actual schema / classes)
# --------------------------------------------------------------------------- #

ERD = """erDiagram
    users ||--o{ orders : places
    users ||--o{ carts : owns
    users ||--o{ content_blocks : edits
    users ||--o{ images : uploads
    event_types ||--o{ events : groups
    venues ||--o{ events : hosts
    restaurants ||--o{ events : hosts
    events ||--o{ ticket_types : offers
    events ||--o{ event_artist : "lines up"
    artists ||--o{ event_artist : performs
    artists ||--o{ artist_images : has
    ticket_types ||--o{ cart_items : "added as"
    ticket_types ||--o{ order_items : "sold as"
    carts ||--o{ cart_items : contains
    orders ||--o{ order_items : contains
    order_items ||--o{ tickets : issues

    users {
        int UserId PK
        varchar Email UK
        varchar FirstName
        varchar LastName
        varchar Role
        varchar Password
        tinyint isVerified
        tinyint isActive
    }
    event_types {
        int id PK
        varchar slug UK
        varchar name
        tinyint is_active
    }
    venues {
        int id PK
        varchar name
        varchar address
        int capacity
    }
    restaurants {
        int id PK
        varchar name
        varchar cuisine
        tinyint stars
        decimal price_per_seat
    }
    artists {
        int id PK
        varchar name
        varchar genre
        text career_highlights
        text tracks
        varchar audio_url
        varchar image
    }
    artist_images {
        int id PK
        int artist_id FK
        varchar path
        int sort_order
    }
    events {
        int id PK
        int event_type_id FK
        int venue_id FK
        int restaurant_id FK
        varchar title
        datetime starts_at
        tinyint is_published
        tinyint is_pass
    }
    event_artist {
        int event_id FK
        int artist_id FK
    }
    ticket_types {
        int id PK
        int event_id FK
        varchar name
        decimal price
        decimal vat_rate
        int capacity
        int sold
        tinyint is_donation
    }
    carts {
        int id PK
        int user_id FK
        varchar session_id
    }
    cart_items {
        int id PK
        int cart_id FK
        int ticket_type_id FK
        int quantity
        varchar special_requests
        decimal custom_price
    }
    orders {
        int id PK
        int user_id FK
        varchar status
        varchar invoice_number UK
        decimal subtotal
        decimal vat_total
        decimal total
        datetime pay_later_until
        datetime paid_at
    }
    order_items {
        int id PK
        int order_id FK
        int ticket_type_id FK
        int quantity
        decimal unit_price
        decimal vat_rate
        varchar special_requests
    }
    tickets {
        int id PK
        int order_item_id FK
        varchar qr_code UK
        varchar status
        datetime scanned_at
    }
    content_blocks {
        int id PK
        varchar page_slug
        varchar block_key
        mediumtext html
        int updated_by FK
    }
    images {
        int id PK
        varchar path
        int uploaded_by FK
    }
"""

DOMAIN = """classDiagram
    class UserModel {
        +int UserId
        +string Email
        +string FirstName
        +string LastName
        +UserRole Role
        +bool isVerified
        +bool isActive
    }
    class EventModel {
        +int id
        +string title
        +string starts_at
        +bool is_published
        +bool is_pass
        +VenueModel venue
        +RestaurantModel restaurant
        +ArtistModel[] artists
    }
    class TicketTypeModel {
        +int id
        +int event_id
        +string name
        +float price
        +float vat_rate
        +int capacity
        +int sold
        +bool is_donation
        +available() int
        +isSoldOut() bool
    }
    class VenueModel {
        +int id
        +string name
        +string address
        +int capacity
    }
    class RestaurantModel {
        +int id
        +string name
        +string cuisine
        +int stars
        +float price_per_seat
    }
    class ArtistModel {
        +int id
        +string name
        +string genre
        +string bio
        +string career_highlights
        +string tracks
        +string audio_url
        +string[] images
        +trackList() string[]
    }
    class CartItemModel {
        +int ticket_type_id
        +int quantity
        +float price
        +float custom_price
        +string special_requests
        +effectivePrice() float
        +lineSubtotal() float
    }
    class OrderModel {
        +int id
        +int user_id
        +string status
        +string invoice_number
        +float subtotal
        +float vat_total
        +float total
        +string pay_later_until
        +OrderItemModel[] items
        +isPaid() bool
        +canPayLater() bool
    }
    class OrderItemModel {
        +int ticket_type_id
        +int quantity
        +float unit_price
        +float vat_rate
        +string special_requests
        +lineTotal() float
    }
    class TicketModel {
        +int id
        +int order_item_id
        +string qr_code
        +string status
    }

    EventModel "1" --> "0..1" VenueModel : at
    EventModel "1" --> "0..1" RestaurantModel : at
    EventModel "*" --> "*" ArtistModel : line-up
    EventModel "1" --> "*" TicketTypeModel : offers
    OrderModel "1" --> "*" OrderItemModel : contains
    OrderModel "*" --> "1" UserModel : placed by
    OrderItemModel "1" --> "1" TicketTypeModel : of
    OrderItemModel "1" --> "*" TicketModel : issues
    CartItemModel "1" --> "1" TicketTypeModel : of
"""

ARCH = """classDiagram
    direction LR
    class CheckoutController {
        +start()
        +success()
        +cancel()
    }
    class CustomerOrderController {
        +index()
        +pay()
    }
    class IOrderService {
        <<interface>>
        +createFromCart(userId) array
        +getByUser(userId) OrderModel[]
        +getByIdForUser(id, userId) OrderModel
        +canStartPayment(order) array
        +getAllForAdmin(status) array
    }
    class OrderService {
        +createFromCart(userId) array
        +canStartPayment(order) array
    }
    class IOrderRepository {
        <<interface>>
        +create(order) int
        +getById(id) OrderModel
        +getByIdForUser(id, userId) OrderModel
        +markPaid(id, invoice)
        +issueTickets(id)
    }
    class OrderRepository {
        +create(order) int
        +issueTickets(id)
    }
    class Repository {
        <<abstract>>
        #fetchOne()
        #fetchAll()
        #execute()
    }
    class OrderModel
    class PaymentService {
        +createCheckoutSession(order, ...) string
    }

    CheckoutController ..> IOrderService : uses
    CustomerOrderController ..> IOrderService : uses
    CheckoutController ..> PaymentService : uses
    IOrderService <|.. OrderService : implements
    OrderService ..> IOrderRepository : uses
    IOrderRepository <|.. OrderRepository : implements
    OrderRepository --|> Repository : extends
    OrderService ..> OrderModel : returns
"""

DIAGRAMS = [
    ("erd", "Entity-Relationship Diagram", ERD),
    ("domain", "Domain Model (UML Class Diagram)", DOMAIN),
    ("architecture", "Application Architecture (Layered MVC)", ARCH),
]


def render(mermaid_src: str) -> bytes:
    b64 = base64.urlsafe_b64encode(mermaid_src.encode()).decode()
    url = "https://mermaid.ink/img/" + b64 + "?type=jpeg&bgColor=ffffff"
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    return urllib.request.urlopen(req, timeout=60).read()


def main() -> None:
    # Render diagrams to JPEG.
    images = {}
    for key, _title, src in DIAGRAMS:
        print("rendering", key, "...", end=" ", flush=True)
        data = render(src)
        path = os.path.join(OUT_DIR, f"diagram_{key}.jpg")
        with open(path, "wb") as fh:
            fh.write(data)
        images[key] = path
        print("ok", len(data), "bytes")

    build_docx(images)
    build_markdown()
    print("done")


# --------------------------------------------------------------------------- #
def build_docx(images: dict) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("Haarlem Festival", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Technical Documentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(16)
    meta = doc.add_paragraph("Ticketing website for The Festival (Inholland, 2.3)\nPlain PHP MVC · MySQL · Docker")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The Haarlem Festival website lets visitors browse six festival events, build a "
        "personal program, buy tickets and passes, make restaurant reservations, donate "
        "to pay-as-you-like events, and pay online (or within 24 hours). Staff scan tickets "
        "at the entrance and administrators manage the catalogue, content and orders."
    )
    doc.add_paragraph(
        "It is built in plain PHP (no framework) following an MVC structure with a "
        "Controller -> Service -> Repository layering. It runs in Docker (nginx, PHP-FPM, "
        "MySQL 8, phpMyAdmin, MailHog) and uses Stripe (test mode) for payment."
    )

    doc.add_heading("2. Architecture", level=1)
    doc.add_paragraph(
        "Requests are routed (FastRoute) to a Controller, which validates input and delegates "
        "to a Service. Services hold the business rules and depend on Repositories for data "
        "access; every Service and Repository is consumed through an interface, so behaviour "
        "is programmed against abstractions, not concrete classes. Repositories extend a base "
        "Repository that wraps PDO with prepared statements."
    )
    doc.add_paragraph(
        "Cross-cutting concerns: AuthMiddleware (authentication, role checks, CSRF), a small "
        "Container for wiring, View for rendering templates, and Flash for one-shot messages. "
        "The diagram below shows the pattern for the Orders slice."
    )
    doc.add_picture(images["architecture"], width=Inches(6.3))
    _caption(doc, "Figure 1 — Layered MVC with interface-based services and repositories (Orders slice).")

    doc.add_heading("3. Database design", level=1)
    doc.add_paragraph(
        "The schema is created and evolved through numbered, forward-only SQL migrations run by "
        "database/migrate.php. Money is stored as DECIMAL; prices are VAT-inclusive and the VAT "
        "portion is derived per line. The entity-relationship diagram below reflects the live schema."
    )
    doc.add_picture(images["erd"], width=Inches(6.5))
    _caption(doc, "Figure 2 — Entity-Relationship Diagram (generated from the live database).")

    doc.add_heading("3.1 Key entities", level=2)
    _entity_table(doc)

    doc.add_heading("4. Domain model", level=1)
    doc.add_paragraph(
        "The domain classes mirror the schema and carry small behaviours. For example "
        "TicketTypeModel::available() derives remaining stock, CartItemModel::effectivePrice() "
        "applies a donation amount or HaarlemPas discount, and OrderModel::canPayLater() encodes "
        "the 24-hour pay-later rule."
    )
    doc.add_picture(images["domain"], width=Inches(6.5))
    _caption(doc, "Figure 3 — Domain model UML class diagram.")

    doc.add_heading("5. Key design decisions", level=1)
    for head, body in [
        ("Program against interfaces",
         "Each Service/Repository has an interface (e.g. IOrderService, IOrderRepository). "
         "Controllers and Services depend on the interface, which keeps layers swappable and testable."),
        ("Migrations",
         "Schema changes are append-only SQL files applied in order and tracked in a migrations "
         "table, so every environment converges to the same state."),
        ("Passes as ticket types",
         "All-access passes are modelled as ticket types on a flagged 'pass event' (events.is_pass), "
         "so they reuse the cart, checkout, ticket and personal-program flow unchanged."),
        ("Effective price",
         "Donations (pay-what-you-like) and the HaarlemPas 25% reduction both resolve to a per-line "
         "effective price (cart_items.custom_price), which becomes the order line's unit_price."),
        ("Reservations",
         "Restaurant reservations charge a EUR 10 per-person fee and capture special requests "
         "(allergies) on the order line, visible to admins on the order detail page."),
        ("Pay later",
         "Orders are created pending with pay_later_until = now + 24h; retry re-validates the "
         "deadline and stock before re-opening Stripe."),
        ("Security & GDPR",
         "Passwords are hashed; all POST requests carry a CSRF token; queries use prepared "
         "statements; the session cookie is HttpOnly/SameSite. Users can export or erase their "
         "data; erasure anonymises the row to preserve invoice integrity."),
    ]:
        p = doc.add_paragraph()
        p.add_run(head + ". ").bold = True
        p.add_run(body)

    out = os.path.join(OUT_DIR, "Technical-Documentation.docx")
    doc.save(out)
    print("wrote", out)


def _caption(doc, text):
    from docx.shared import Pt
    p = doc.add_paragraph(text)
    r = p.runs[0]
    r.italic = True
    r.font.size = Pt(9)


def _entity_table(doc):
    rows = [
        ("users", "Accounts with a role (customer/employee/admin)."),
        ("event_types", "The six festival events (Jazz, DANCE!, Yummy, History, Magic, Stories)."),
        ("events", "A session/performance under a type, at a venue or restaurant; may be a pass."),
        ("ticket_types", "What can be bought for an event: price, VAT, capacity, donation flag."),
        ("venues / restaurants / artists", "Catalogue entities; artists link to events many-to-many and have detail content."),
        ("artist_images", "Gallery images for an artist's detail page."),
        ("carts / cart_items", "A guest or user cart; lines hold quantity, requests and custom price."),
        ("orders / order_items", "A placed order and its priced lines; status drives pay-later."),
        ("tickets", "Issued per order-item with a unique QR code; scanned at the entrance."),
        ("content_blocks / images", "CMS content (WYSIWYG HTML) and uploaded images."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].add_run("Entity").bold = True
    hdr[1].paragraphs[0].add_run("Purpose").bold = True
    for name, desc in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = desc


# --------------------------------------------------------------------------- #
def build_markdown() -> None:
    md = f"""# Haarlem Festival — Technical Documentation

Ticketing website for The Festival (Inholland 2.3). Plain PHP MVC, MySQL, Docker,
Stripe (test) payment.

## 1. Architecture

Requests are routed to a **Controller**, which delegates to a **Service** (business
rules) that uses a **Repository** (data access). Every Service and Repository is
consumed through an **interface**, so the code is programmed against abstractions.
Repositories extend a base `Repository` wrapping PDO with prepared statements.
Cross-cutting: `AuthMiddleware` (auth, roles, CSRF), `View`, `Flash`, `Container`.

```mermaid
{ARCH}```

## 2. Database design

The schema is built by numbered, forward-only SQL **migrations**
(`database/migrate.php`). Money is `DECIMAL`; prices are VAT-inclusive and the VAT
portion is derived per line.

```mermaid
{ERD}```

## 3. Domain model

Domain classes mirror the schema and carry small behaviours — e.g.
`TicketTypeModel::available()`, `CartItemModel::effectivePrice()` (donation /
HaarlemPas), `OrderModel::canPayLater()` (24-hour rule).

```mermaid
{DOMAIN}```

## 4. Key design decisions

- **Program against interfaces** — `IOrderService`, `IOrderRepository`, etc.
- **Migrations** — append-only, tracked, reproducible.
- **Passes as ticket types** — on a flagged `is_pass` event, reusing the whole flow.
- **Effective price** — donations and the HaarlemPas 25% both resolve to
  `cart_items.custom_price`, which becomes the order line `unit_price`.
- **Reservations** — EUR 10 per-person fee + special requests on the order line.
- **Pay later** — `pay_later_until = now + 24h`; retry re-validates deadline + stock.
- **Security & GDPR** — hashed passwords, CSRF on every POST, prepared statements,
  hardened session cookie, data export + anonymising erasure.
"""
    out = os.path.join(OUT_DIR, "technical-documentation.md")
    with open(out, "w", encoding="utf-8") as fh:
        fh.write(md)
    print("wrote", out)


if __name__ == "__main__":
    main()
